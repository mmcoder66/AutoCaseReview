"""Generate the SVN Word deliverable (LC-SOP-RC-003-M01).

Strategy: open the template, fill cells inside the only table (``会议纪要``),
dynamically extend the ``改进或遗留工作项`` section with one row per todo
item.  All template styles, fonts and merged cells are preserved.
"""

from __future__ import annotations

import argparse
from copy import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from data_loader import (
    OUTPUTS_DIR,
    TEMPLATES_DIR,
    derive_meeting_name,
    ensure_outputs_dir,
    expand_todos,
    get_default_product,
    get_output_subdir,
    list_iterations,
    load_content_rules,
    load_filename_templates,
    load_template_names,
    render_filename,
    resolve_date_by_strategy,
)

# Row indices inside doc.tables[0].
ROW_MEETING_NAME = 1
ROW_MEETING_PLACE = 2
ROW_MEETING_TIME = 3
ROW_PARTICIPANTS = 4
ROW_CONTENT_HEADER = 7
ROW_CONTENT_BODY = 8
ROW_TODO_HEADER = 10
FIRST_TODO_ROW = 11   # template ships rows 11 and 12 as samples.

# Column indices inside the table.
COL_LABEL = 0
COL_MEETING_NAME_VAL = 2
COL_MEETING_TIME_VAL = 2
COL_RECORDER_LABEL = 5
COL_RECORDER_VAL = 7
COL_PARTICIPANTS_VAL = 2
COL_CONTENT_BODY = 0       # merged across all 8 cols

# 改进或遗留工作项 columns (after merge): 序号 / 问题描述 / 责任人 / 计划解决日期 / 状态 / 备注
TODO_COL_SEQ = 0
TODO_COL_DESC = 1          # merged with col 2
TODO_COL_OWNER = 3
TODO_COL_DATE = 4          # merged with col 5
TODO_COL_STATUS = 6
TODO_COL_NOTE = 7


# ---------------------------------------------------------------------------
# Cell helpers that preserve formatting
# ---------------------------------------------------------------------------
def _set_cell_text(cell, text: str) -> None:
    """Replace a cell's text while keeping its paragraph / run style.

    Strategy: reuse the first paragraph and first run when present so that
    fonts, sizes and alignment survive the update.  Extra paragraphs inside
    the cell are cleared.
    """
    paragraphs = cell.paragraphs
    target_para = paragraphs[0]

    # Drop trailing paragraphs (keeps cell clean for multi-line content later).
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # Reuse existing run if any, otherwise create one.
    runs = target_para.runs
    if runs:
        first_run = runs[0]
        first_run.text = str(text)
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        target_para.text = str(text)


def _set_multiline_cell(cell, lines: list[str]) -> None:
    """Write *lines* as separate paragraphs inside *cell* (style preserved)."""
    if not lines:
        _set_cell_text(cell, "")
        return
    paragraphs = cell.paragraphs
    target_para = paragraphs[0]
    # Capture style of the first run if any so we can clone it.
    template_run = target_para.runs[0] if target_para.runs else None

    # Remove all paragraphs first.
    for p in paragraphs:
        p._element.getparent().remove(p._element)

    for line in lines:
        new_para = cell.add_paragraph()
        if template_run is not None:
            new_run = new_para.add_run(str(line))
            new_run.font.name = template_run.font.name
            new_run.font.size = template_run.font.size
            new_run.font.bold = template_run.font.bold
            rPr = new_run._element.get_or_add_rPr()
            rFonts = template_run._element.rPr.rFonts if template_run._element.rPr is not None else None
            if rFonts is not None:
                new_rFonts = rPr.find(qn("w:rFonts"))
                if new_rFonts is None:
                    from docx.oxml import OxmlElement
                    new_rFonts = OxmlElement("w:rFonts")
                    rPr.append(new_rFonts)
                for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
                    if rFonts.get(qn(attr)):
                        new_rFonts.set(qn(attr), rFonts.get(qn(attr)))
        else:
            new_para.add_run(str(line))


def _clone_row(template_row, table):
    """Append a new row that mirrors *template_row*'s cell grid + merges."""
    new_row = table.add_row()
    # Copy per-cell text/style would be overkill; we just mirror merge grid.
    return new_row


def _apply_merge_grid(row, merge_pairs: list[tuple[int, int]]) -> None:
    """Merge the requested column pairs inside *row*."""
    for start, end in merge_pairs:
        a = row.cells[start]
        b = row.cells[end]
        a.merge(b)


# ---------------------------------------------------------------------------
# Section fillers
# ---------------------------------------------------------------------------
def _fill_meeting_metadata(table, *, meeting_name, meeting_place, meeting_time,
                           recorder, participants) -> None:
    _set_cell_text(table.rows[ROW_MEETING_NAME].cells[COL_MEETING_NAME_VAL], meeting_name)
    _set_cell_text(table.rows[ROW_MEETING_PLACE].cells[COL_MEETING_NAME_VAL], meeting_place)
    _set_cell_text(table.rows[ROW_MEETING_TIME].cells[COL_MEETING_TIME_VAL], meeting_time)
    _set_cell_text(table.rows[ROW_MEETING_TIME].cells[COL_RECORDER_VAL], recorder)
    _set_cell_text(table.rows[ROW_PARTICIPANTS].cells[COL_PARTICIPANTS_VAL], participants)


def _build_meeting_content(df) -> str:
    """Fill ``会议内容`` with all todo text from requirement_data.

    SVN Word is generated per requirement, so this writes every
    ``代办事项N@责任人`` value from the current requirement.  The source text is
    preserved, including @mentions.
    """
    content_cfg = load_content_rules().get("meeting", {}).get("content", {})
    todos = expand_todos(df)
    if not todos:
        return str(content_cfg.get("no_todos_text", "无"))
    return "\n".join(str(item["description"]).strip() for item in todos if item["description"])


def _fill_meeting_content(table, df) -> None:
    body_cell = table.rows[ROW_CONTENT_BODY].cells[COL_CONTENT_BODY]
    content = _build_meeting_content(df)
    if not content:
        _set_cell_text(body_cell, "")
        return
    _set_multiline_cell(body_cell, content.split("\n"))


def _clear_existing_todo_rows(table) -> None:
    """Remove the template's placeholder rows 11+ from the todo section."""
    while len(table.rows) > FIRST_TODO_ROW:
        last = table.rows[-1]
        last._element.getparent().remove(last._element)


def _append_todo_row(table, item: dict, todo_cfg: dict) -> None:
    """Append one row to the ``改进或遗留工作项`` section with proper merges.

    All presentation fields come from *todo_cfg* (= ``content_rules.yaml::todo``):
    * ``plan_date`` — strategy-driven (default: day before 计划完成日期)
    * ``status``    — fixed string (default: "已完成")
    * ``note``      — fixed string (default: "")
    """
    row = table.add_row()
    # Mirror the template's merge grid: cols 1-2 (问题描述) and 4-5 (计划解决日期).
    _apply_merge_grid(row, [(TODO_COL_DESC, TODO_COL_DESC + 1),
                            (TODO_COL_DATE, TODO_COL_DATE + 1)])

    owners = "、".join(item["owners"]) if item["owners"] else ""
    # Build a tiny dict so resolve_date_by_strategy can read 计划完成日期 from the item.
    date_source = {"计划完成日期": item.get("plan_date_raw", "")}
    plan_date = resolve_date_by_strategy(todo_cfg.get("plan_date", {}), date_source)
    status = todo_cfg.get("status", "")
    note = todo_cfg.get("note", "")

    _set_cell_text(row.cells[TODO_COL_SEQ], str(item["seq"]))
    _set_cell_text(row.cells[TODO_COL_DESC], item["description"])
    _set_cell_text(row.cells[TODO_COL_OWNER], owners)
    _set_cell_text(row.cells[TODO_COL_DATE], plan_date)
    _set_cell_text(row.cells[TODO_COL_STATUS], status)
    _set_cell_text(row.cells[TODO_COL_NOTE], note)


def _append_no_todo_row(table, todo_cfg: dict) -> None:
    """Append one visible placeholder row when there are no todo items."""
    row = table.add_row()
    _apply_merge_grid(row, [(TODO_COL_DESC, TODO_COL_DESC + 1),
                            (TODO_COL_DATE, TODO_COL_DATE + 1)])
    _set_cell_text(row.cells[TODO_COL_SEQ], "1")
    _set_cell_text(row.cells[TODO_COL_DESC], str(todo_cfg.get("no_items_text", "无")))
    _set_cell_text(row.cells[TODO_COL_OWNER], "")
    _set_cell_text(row.cells[TODO_COL_DATE], "")
    _set_cell_text(row.cells[TODO_COL_STATUS], "")
    _set_cell_text(row.cells[TODO_COL_NOTE], "")


def _fill_todo_section(table, df) -> int:
    _clear_existing_todo_rows(table)
    todo_cfg = load_content_rules().get("todo", {})
    todos = expand_todos(df)
    if not todos:
        _append_no_todo_row(table, todo_cfg)
        return 0
    for item in todos:
        _append_todo_row(table, item, todo_cfg)
    return len(todos)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def generate(
    df,
    *,
    product: str | None = None,
    version: str = "A1",
    meeting_name: str | None = None,
    meeting_place: str = "线上会议",
    meeting_time: str = "",
    recorder: str = "",
    participants: str | None = None,
    subject: str = "",
    output_name: str | None = None,
    output_dir: Path | str | None = None,
) -> Path:
    product_value = product or get_default_product()
    template_path = TEMPLATES_DIR / load_template_names()["svn_word"]
    if not template_path.exists():
        raise FileNotFoundError(f"template missing: {template_path}")

    doc = Document(str(template_path))
    if not doc.tables:
        raise RuntimeError("template has no table; cannot fill")
    table = doc.tables[0]

    resolved_name = meeting_name or derive_meeting_name(df, product=product_value)
    _fill_meeting_metadata(
        table,
        meeting_name=resolved_name,
        meeting_place=meeting_place,
        meeting_time=meeting_time,
        recorder=recorder,
        participants=participants or "",
    )
    _fill_meeting_content(table, df)
    _fill_todo_section(table, df)

    # Filename: render from YAML template (see config/filename_templates.yaml).
    iterations = list_iterations(df)
    if len(iterations) == 1:
        iteration_value = iterations[0]
    elif iterations:
        iteration_value = "多迭代"
    else:
        iteration_value = ""
    req_id_value = ""
    if len(df) == 1:
        req_id_value = str(df.iloc[0].get("ID", "")).strip().lstrip("#")

    templates = load_filename_templates()
    name = output_name or render_filename(
        templates["svn_word"],
        product=product_value,
        title=subject,
        iteration=iteration_value,
        version=version or "",
        req_id=req_id_value,
    )

    subdir = get_output_subdir("svn_word", templates)
    base_dir = Path(output_dir) if output_dir else ensure_outputs_dir()
    out_dir = base_dir / subdir if subdir else base_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / name
    if out_path.exists():
        out_path.unlink()
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate SVN Word (LC-SOP-RC-003-M01).")
    parser.add_argument("--iteration", help="Filter on 所属迭代 before generating.")
    parser.add_argument("--product", default=get_default_product())
    parser.add_argument("--version", default="A1")
    parser.add_argument("--meeting-name", default=None)
    parser.add_argument("--meeting-place", default="线上会议")
    parser.add_argument("--meeting-time", default="")
    parser.add_argument("--recorder", default="")
    parser.add_argument("--participants", default=None,
                        help="Comma-separated; auto-derived from data if omitted.")
    parser.add_argument("--data-dir", default=None)
    parser.add_argument("--output-dir", default=None)
    parser.add_argument("--output-name", default=None)
    args = parser.parse_args()

    from data_loader import collect_participants, load_all_requirements

    df = load_all_requirements(args.data_dir, iteration=args.iteration)
    participants = args.participants
    if participants is None:
        participants = "、".join(collect_participants(df))

    out = generate(
        df,
        product=args.product,
        version=args.version,
        meeting_name=args.meeting_name,
        meeting_place=args.meeting_place,
        meeting_time=args.meeting_time,
        recorder=args.recorder,
        participants=participants,
        output_name=args.output_name,
        output_dir=args.output_dir,
    )
    print(f"OK  rows={len(df)}  -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
