"""Generate the test-case-review email Word (single iteration).

Layout mirrors ``inputs/templates/email.png``:

  <configured greeting>
  <configured intro>

  <configured review method>

  <configured requirements title>
  ┌──────┬──────────────────────────────────┬──────┬──────────┬──────┐
  │ 编号 │ 标题                              │ 测试 │ 评审结果 │ 备注# │
  ├──────┼──────────────────────────────────┼──────┼──────────┼──────┤
  │ ...  │ ...                              │ ...  │ 通过     │      │
  └──────┴──────────────────────────────────┴──────┴──────────┴──────┘

  3.评审记录/代办事项
  #需求ID 需求标题
  • 代办事项原文（含 @责任人）
  ...
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from data_loader import (
    OUTPUTS_DIR,
    ensure_outputs_dir,
    expand_todos,
    get_default_product,
    get_output_subdir,
    list_iterations,
    load_content_rules,
    load_filename_templates,
    render_filename,
)

DEFAULT_FONT = "微软雅黑"
DEFAULT_SIZE_PT = 10.5
SECTION_TITLE_SIZE_PT = 12


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def _apply_chinese_font(
    run,
    font_name: str = DEFAULT_FONT,
    size_pt: float = DEFAULT_SIZE_PT,
) -> None:
    """Make a run render Chinese characters correctly under python-docx."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font_name)


def _set_cell_border(cell) -> None:
    """Add a thin single-line border on all four sides of *cell*."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        border = tcBorders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tcBorders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")    # 0.5pt
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _add_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    indent: float = 0.0,
    size_pt: float = DEFAULT_SIZE_PT,
) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(indent * 10)
    run = p.add_run(text)
    run.bold = bold
    _apply_chinese_font(run, size_pt=size_pt)
    return p


def _add_header_cell_shading(cell, fill: str = "D9E1F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _fill_header_row(row, headers: list[str]) -> None:
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Replace default empty paragraph
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        _apply_chinese_font(run)
        _add_header_cell_shading(cell)
        _set_cell_border(cell)


def _fill_body_row(row, requirement, *, review_result: str) -> None:
    rid = str(requirement.get("ID", "")).strip()
    title = str(requirement.get("标题", "")).strip()
    tester = str(requirement.get("测试", "")).strip()
    values = [rid, title, tester, review_result, ""]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(value)
        _apply_chinese_font(run)
        _set_cell_border(cell)


def _build_review_table(doc, df, *, headers: list[str], review_result: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    _fill_header_row(table.rows[0], headers)
    for _, row in df.iterrows():
        body_row = table.add_row()
        _fill_body_row(body_row, row, review_result=review_result)


def _build_intro_line(template: str, product: str, iteration: str) -> str:
    return template.format(product=product, iteration=iteration)


def _build_todo_section(doc, df) -> int:
    """Section 3: ``评审记录/代办事项`` — group todos under each requirement."""
    emitted = 0
    for _, row in df.iterrows():
        todos = expand_todos(row.to_frame().T)
        if not todos:
            continue
        title = str(row.get("标题", "")).strip()
        rid = str(row.get("ID", "")).strip()
        _add_paragraph(doc, f"{rid} {title}".strip())
        for todo in todos:
            _add_paragraph(doc, f"• {todo['description']}", indent=0.5)
            emitted += 1
    return emitted


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def generate(
    df,
    *,
    product: str | None = None,
    iteration: str | None = None,
    output_name: str | None = None,
    output_dir: Path | str | None = None,
) -> Path:
    if len(list_iterations(df)) > 1:
        raise ValueError(
            "Email deliverable supports a single iteration only. "
            "Filter the DataFrame on 所属迭代 before calling generate()."
        )

    doc = Document()
    # Narrower margins so wide tables fit on screen-friendly pages.
    section = doc.sections[0]
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    product_value = product or get_default_product()
    email_cfg = load_content_rules().get("email", {})
    iterations = list_iterations(df)
    iteration_value = iteration or (iterations[0] if iterations else "评审")

    # Greeting + section 1.  Keep the wording aligned with inputs/templates/email.png.
    _add_paragraph(doc, email_cfg["greeting"])
    _add_paragraph(
        doc,
        _build_intro_line(
            email_cfg["intro_template"],
            product_value,
            iteration_value,
        ),
        indent=2.5,
    )
    _add_paragraph(doc, email_cfg["review_method"], bold=True, size_pt=SECTION_TITLE_SIZE_PT)
    # Section 2 header (table follows immediately).
    _add_paragraph(doc, email_cfg["requirements_title"], bold=True, size_pt=SECTION_TITLE_SIZE_PT)

    _build_review_table(
        doc,
        df,
        headers=list(email_cfg["table_headers"]),
        review_result=email_cfg["review_result"],
    )

    _add_paragraph(doc, "")  # spacer
    _add_paragraph(doc, email_cfg["todos_title"], bold=True, size_pt=SECTION_TITLE_SIZE_PT)
    if _build_todo_section(doc, df) == 0:
        _add_paragraph(doc, email_cfg["no_todos_text"], indent=0.5)

    templates = load_filename_templates()
    name = output_name or render_filename(
        templates["email_word"],
        product=product_value,
        iteration=iteration_value,
        version="",  # email never carries a version tag
    )

    subdir = get_output_subdir("email_word", templates)
    base_dir = Path(output_dir) if output_dir else ensure_outputs_dir()
    out_dir = base_dir / subdir if subdir else base_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / name
    if out_path.exists():
        out_path.unlink()
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate email Word for one iteration.")
    parser.add_argument("--iteration", required=True,
                        help="所属迭代 value; email deliverable is single-iteration only.")
    parser.add_argument("--product", default=get_default_product())
    parser.add_argument("--data-dir", default=None)
    parser.add_argument("--output-dir", default=None)
    parser.add_argument("--output-name", default=None)
    args = parser.parse_args()

    from data_loader import load_all_requirements

    df = load_all_requirements(args.data_dir, iteration=args.iteration)
    if df.empty:
        raise SystemExit(f"no rows for iteration {args.iteration!r}")

    out = generate(
        df,
        product=args.product,
        output_name=args.output_name,
        output_dir=args.output_dir,
    )
    print(f"OK  iteration={args.iteration} rows={len(df)} -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
